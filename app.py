import streamlit as st
from PyPDF2 import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from groq import Groq
from tavily import TavilyClient
import docx
import openpyxl

st.set_page_config(page_title="AI Assistant", page_icon="🤖", layout="centered")

groq_key = st.secrets["GROQ_API_KEY"]
tavily_key = st.secrets["TAVILY_API_KEY"]
client = Groq(api_key=groq_key)
tavily = TavilyClient(api_key=tavily_key)

st.markdown("""
<style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stApp {
        background: linear-gradient(135deg, #0f0c29, #302b63, #24243e);
        max-width: 100% !important;
    }
    [data-testid="stSidebar"] {
        background: rgba(255,255,255,0.05) !important;
        border-right: 1px solid rgba(255,255,255,0.1);
    }
    [data-testid="stFileUploader"] {
        background: rgba(255,255,255,0.05) !important;
        border: 2px dashed rgba(255,255,255,0.3) !important;
        border-radius: 10px !important;
    }
    [data-testid="stFileUploader"] * {
        color: white !important;
        background: transparent !important;
    }
    [data-testid="stFileUploaderDropzone"] {
        background: rgba(255,255,255,0.05) !important;
    }
    [data-testid="stFileUploaderDropzone"] button {
        background: rgba(102, 126, 234, 0.3) !important;
        color: white !important;
        border: 1px solid rgba(102,126,234,0.5) !important;
        border-radius: 8px !important;
    }
    .stButton button {
        background: linear-gradient(135deg, #667eea, #764ba2) !important;
        color: white !important;
        border: none !important;
        border-radius: 10px !important;
        font-weight: 600 !important;
        padding: 0.6rem !important;
    }
    /* Mode cards */
    .mode-card {
        background: rgba(255,255,255,0.05);
        border: 2px solid rgba(255,255,255,0.1);
        border-radius: 16px;
        padding: 1.2rem;
        text-align: center;
        cursor: pointer;
        transition: all 0.3s ease;
        margin: 0.3rem;
    }
    .mode-card:hover {
        background: rgba(102,126,234,0.2);
        border-color: rgba(102,126,234,0.5);
        transform: translateY(-2px);
    }
    .mode-card.active {
        background: linear-gradient(135deg, rgba(102,126,234,0.3), rgba(118,75,162,0.3));
        border-color: #667eea;
        box-shadow: 0 0 20px rgba(102,126,234,0.3);
    }
    .mode-icon { font-size: 2rem; margin-bottom: 0.5rem; }
    .mode-title { font-size: 0.95rem; font-weight: 700; color: white !important; margin-bottom: 0.3rem; }
    .mode-desc { font-size: 0.75rem; color: rgba(255,255,255,0.5) !important; }
    /* Chat */
    [data-testid="stChatMessage"] {
        padding: 0.5rem !important;
        border-radius: 12px !important;
        max-width: 100% !important;
    }
    h1, h2, h3, p, span, label, div { color: white !important; }
    hr { border-color: rgba(255,255,255,0.1) !important; }
    @media (max-width: 768px) {
        h1 { font-size: 1.5rem !important; }
        .mode-icon { font-size: 1.5rem; }
        .mode-title { font-size: 0.85rem; }
    }
</style>
""", unsafe_allow_html=True)

def extract_text(uploaded_file):
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        pdf_reader = PdfReader(uploaded_file)
        text = ""
        for page in pdf_reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text
    elif name.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        text = ""
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text.strip() + "\n"
        return text
    elif name.endswith(".xlsx"):
        wb = openpyxl.load_workbook(uploaded_file)
        text = ""
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            text += f"Sheet: {sheet}\n"
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join([str(cell) for cell in row if cell is not None])
                if row_text:
                    text += row_text + "\n"
        return text
    elif name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")
    return ""

def search_internet(query):
    try:
        results = tavily.search(query=query, max_results=5)
        context = ""
        for r in results.get("results", []):
            context += f"Title: {r.get('title','')}\nContent: {r.get('content','')}\nURL: {r.get('url','')}\n\n"
        return context if context else "No results found."
    except Exception as e:
        return f"Search error: {str(e)}"

def get_history():
    return [{"role": m["role"], "content": m["content"]} for m in st.session_state.chat_history[-10:]]

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "doc_text" not in st.session_state:
    st.session_state.doc_text = ""
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "doc_processed" not in st.session_state:
    st.session_state.doc_processed = False
if "doc_name" not in st.session_state:
    st.session_state.doc_name = ""
if "mode" not in st.session_state:
    st.session_state.mode = "document"

# ── Sidebar ───────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 🤖 AI Assistant")
    st.markdown("---")

    # Mode selector cards
    st.markdown("### Choose Mode")
    col1, col2 = st.columns(2)

    with col1:
        doc_active = "active" if st.session_state.mode == "document" else ""
        st.markdown(f"""
        <div class="mode-card {doc_active}" id="doc-card">
            <div class="mode-icon">📄</div>
            <div class="mode-title">Document</div>
            <div class="mode-desc">Chat with your files</div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Select", key="doc_btn", use_container_width=True):
            st.session_state.mode = "document"
            st.rerun()

    with col2:
        web_active = "active" if st.session_state.mode == "internet" else ""
        st.markdown(f"""
        <div class="mode-card {web_active}" id="web-card">
            <div class="mode-icon">🌐</div>
            <div class="mode-title">Internet</div>
            <div class="mode-desc">Search the web</div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Select", key="web_btn", use_container_width=True):
            st.session_state.mode = "internet"
            st.rerun()

    st.markdown("---")

    if st.session_state.mode == "document":
        uploaded_file = st.file_uploader("📎 Upload Document", type=["pdf", "docx", "xlsx", "txt"])
        if st.button("🚀 Process Document", use_container_width=True):
            if not uploaded_file:
                st.error("❌ Upload a file first")
            else:
                with st.spinner("🔍 Reading..."):
                    raw_text = extract_text(uploaded_file)
                    if not raw_text.strip():
                        st.error("❌ Could not extract text.")
                    else:
                        word_count = len(raw_text.split())
                        if word_count <= 4000:
                            st.session_state.doc_text = raw_text
                            st.session_state.vectorstore = None
                        else:
                            splitter = RecursiveCharacterTextSplitter(
                                chunk_size=800, chunk_overlap=150,
                                separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
                            )
                            chunks = splitter.split_text(raw_text)
                            embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
                            st.session_state.vectorstore = FAISS.from_texts(chunks, embeddings)
                            st.session_state.doc_text = ""
                        st.session_state.doc_processed = True
                        st.session_state.doc_name = uploaded_file.name
                        st.success(f"✅ Ready! ({word_count} words)")

        st.markdown("---")
        if st.session_state.doc_processed:
            st.markdown(f"📄 **{st.session_state.doc_name}**")
            st.markdown('<span style="color:#34d399">● Document loaded</span>', unsafe_allow_html=True)
        else:
            st.markdown('<span style="color:#fbbf24">● No document loaded</span>', unsafe_allow_html=True)
    else:
        st.markdown("""
        <div style="background:rgba(255,255,255,0.05);border-radius:12px;padding:1rem;text-align:center">
            <div style="font-size:1.5rem">🌐</div>
            <div style="font-weight:600;margin:0.3rem 0">Internet Search Active</div>
            <div style="font-size:0.8rem;color:rgba(255,255,255,0.5)">Ask me anything — I'll search the web for real time answers</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")
    if st.button("🗑️ Clear Chat", use_container_width=True):
        st.session_state.chat_history = []
        st.rerun()

# ── Main Content ──────────────────────────────────────────────
st.title("🤖 AI Assistant")

if st.session_state.mode == "document":
    if not st.session_state.doc_processed:
        st.markdown("Upload any document and have a full conversation with it")
        st.markdown("---")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown('<div style="background:rgba(255,255,255,0.05);border-radius:10px;padding:0.8rem;text-align:center"><div style="font-size:1.5rem">📄</div><div style="font-size:0.8rem;margin-top:0.3rem">PDF</div></div>', unsafe_allow_html=True)
        with col2:
            st.markdown('<div style="background:rgba(255,255,255,0.05);border-radius:10px;padding:0.8rem;text-align:center"><div style="font-size:1.5rem">📝</div><div style="font-size:0.8rem;margin-top:0.3rem">Word</div></div>', unsafe_allow_html=True)
        with col3:
            st.markdown('<div style="background:rgba(255,255,255,0.05);border-radius:10px;padding:0.8rem;text-align:center"><div style="font-size:1.5rem">📊</div><div style="font-size:0.8rem;margin-top:0.3rem">Excel</div></div>', unsafe_allow_html=True)
        with col4:
            st.markdown('<div style="background:rgba(255,255,255,0.05);border-radius:10px;padding:0.8rem;text-align:center"><div style="font-size:1.5rem">📃</div><div style="font-size:0.8rem;margin-top:0.3rem">Text</div></div>', unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)
        st.warning("👈 Upload your document in the sidebar to get started")
    else:
        st.markdown(f"Chatting with: **{st.session_state.doc_name}**")
else:
    st.markdown("Ask me anything — searching the internet for real time answers!")

st.markdown("---")

for message in st.session_state.chat_history:
    with st.chat_message(message["role"]):
        st.write(message["content"])

placeholder = "Ask about your document..." if st.session_state.mode == "document" else "Ask me anything..."

if prompt := st.chat_input(placeholder):
    st.session_state.chat_history.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.write(prompt)

    with st.chat_message("assistant"):
        with st.spinner("Thinking..."):

            if st.session_state.mode == "internet":
                search_results = search_internet(prompt)
                messages = [
                    {"role": "system", "content": f"""You are a helpful AI assistant with real time internet access.

Search results:
{search_results}

Rules:
- Answer directly and confidently from search results
- Never say your data is limited or has a cutoff
- Never tell user to check other websites
- Remember our conversation history
- Be conversational and natural"""},
                ] + get_history()

            elif st.session_state.doc_text:
                messages = [
                    {"role": "system", "content": f"""You are an expert assistant helping answer questions based on this document.

Complete document:
{st.session_state.doc_text}

Rules:
- Answer in first person as the candidate speaking
- Sound confident natural and conversational
- Tell a story not a list
- Use specific real details from the document
- Remember our conversation history
- Never sound robotic"""},
                ] + get_history()

            elif st.session_state.vectorstore:
                docs = st.session_state.vectorstore.similarity_search(prompt, k=6)
                context = "\n\n".join([doc.page_content for doc in docs])
                messages = [
                    {"role": "system", "content": f"""You are a helpful assistant answering from this document:

{context}

Rules:
- Answer clearly and accurately
- Be conversational and natural
- Remember our conversation history"""},
                ] + get_history()

            else:
                messages = [
                    {"role": "system", "content": "You are a helpful friendly AI assistant. Remember conversation history and answer naturally."},
                ] + get_history()

            response = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=messages
            )
            answer = response.choices[0].message.content
            st.write(answer)
            st.session_state.chat_history.append({"role": "assistant", "content": answer})
